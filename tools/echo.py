# -*- coding: utf-8 -*-
"""
Sowads Echo — feature de plataforma: o líder "ecoa" o conteúdo publicado da empresa na
própria voz. Transforma um artigo do blog em post autoral de LinkedIn (emulação de persona)
e gera um DOCX com 1 post por página. Um post por artigo, um prompt por artigo.

Config de persona atual: Celso Sato (CEO Accesstage) — 1ª pessoa, gancho variado (9 ângulos
rotacionados), insight de liderança, técnico traduzido em benefício humano, fecho com
pergunta, link + hashtags. Sem emoji, sem travessão, sem cacoetes de IA (limpeza pós-processo).
Modelo padrão: anthropic/claude-opus-4.7.

Uso:
  python3 tools/echo.py --csv output/articles/<batch>.csv \
      --url_dir output/accesstage-site/<lote> \
      --url_base https://sowads-accesstage.web.app/<lote> \
      --out output/echo/sowads_echo_<lote>.docx [--limit N] [--model <id>]
"""
import os, re, csv, glob, html, json, argparse, time
import requests
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_BREAK

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
load_dotenv(os.path.join(BASE, ".env"))
KEY = os.getenv("OPENROUTER_API_KEY")
URL = "https://openrouter.ai/api/v1/chat/completions"
MODEL = "anthropic/claude-opus-4.7"

SYSTEM = """Você é Celso Sato, CEO do Grupo Accesstage (techfin líder em Open Finance B2B na América Latina). Sua tarefa: pegar o TEMA de um artigo do blog da Accesstage e escrever um post autoral SEU no LinkedIn — com opinião, energia e voz de dono.

NÃO resuma o artigo. Use o tema como trampolim para a SUA visão de líder. O leitor tem que sentir uma pessoa real, com convicção, não um comunicado.

VOZ (obrigatória):
- Primeira pessoa e "nós". Caloroso, otimista, grato, mas com PONTO DE VISTA claro e firme.
- Comece por algo concreto e vivo (uma cena, uma fala de cliente, um número, uma tese) — nunca por conceito genérico. Só use cena, fala, número ou caso específico quando ele estiver explicitamente no artigo. Caso contrário, trate-o como hipótese geral, sem inventar fonte, cliente ou estatística.
- Extraia UM princípio de liderança memorável (uma frase que gruda).
- Traduza o técnico (conciliação, risco sacado, capital de giro, IA, Open Finance) em benefício humano: tempo, previsibilidade, sono tranquilo, menos dor.
- Credite pessoas quando fizer sentido (time "AccessLovers", clientes).
- Feche com pergunta aberta OU convite de propósito.

ENERGIA / ANTI-TEMPLATE (crítico — o problema a evitar):
- PROIBIDO abrir com "Comecei a manhã lendo", "Comecei o dia", "Nesta manhã", "Nesta semana lendo", ou dizer que leu um artigo. Vá direto ao ângulo.
- PROIBIDO clichê corporativo morto: "é crucial", "é fundamental", "indispensável", "cada vez mais", "no cenário atual", "aliada indispensável", "peça-chave", "de suma importância", "não é apenas sobre X, mas sobre Y".
- Escreva com sangue: uma opinião firme, um detalhe concreto, uma frase curta de impacto. Varie o ritmo (frases curtas e médias). 1–3 exclamações no máximo, quando couber de verdade.
- Cada post tem que soar DIFERENTE do anterior — abertura, ritmo e ângulo variados.

HUMANIDADE (tem que parecer escrito por uma PESSOA, não por IA):
- PROIBIDO travessão (— ou –). Nunca use travessão para pausa ou aposto. Use vírgula, ponto ou reticências.
- PROIBIDO ponto e vírgula.
- Sem cacoetes de IA: nada de "não é X, é Y" / "não é sobre X, mas sobre Y"; nada de listas de três toda hora; nada de scaffolding tipo "Minha tese é:", "Aqui vai minha tese:", "O que muda na prática:", "Minha convicção é esta:", "Isso não é teoria. É...", "Simples. Direta. Verdadeira."
- Escreva fluido e humano, com imperfeição natural. Nada de parágrafos milimetricamente paralelos. Português-BR impecável.

REGRAS: nunca copie frases literais do artigo. Nunca soe como release ("A Accesstage anuncia"). Sem venda agressiva, sem polêmica. Sem emojis, salvo se a instrução do usuário liberar explicitamente um único emoji contextual. Parágrafos curtos (1–3 frases), com linha em branco entre eles. 600–1300 caracteres. Inclua o link do artigo, mas NÃO inclua hashtags: elas são selecionadas e validadas pelo sistema.

Bordões que pode usar com naturalidade (sem empilhar): "colocar as pessoas no centro", "construindo, juntos", "empresa longeva", "crescimento sustentável", "resultados através de pessoas felizes", "vestir a camisa", "Quer ir rápido, vá sozinho; quer ir longe, vá com um time".

Hashtags disponíveis (escolha 3–5 por tema): #Accesstage #GrupoAccesstage #GestãoFinanceira #Tesouraria #ClienteNoCentro #RiscoSacado #AntecipacaoDeRecebiveis #CapitalDeGiro #TecnologiaFinanceira #CrescimentoSustentavel #MercadoFinanceiro #Inovacao #OpenFinance

OUTPUT: apenas o texto final do post, pronto para publicar (com link e sem hashtags). Sem rótulos, sem aspas ao redor."""

# Ângulos de abertura rotacionados (garante variação real entre posts)
OPENING_ANGLES = [
    "TESE PROVOCATIVA: abra com uma afirmação forte e contra-intuitiva sobre o tema (ex.: 'Tesouraria não é sobre fechar o caixa.').",
    "FALA DE CLIENTE: abra com uma cena/fala real de um CFO ou gestor que você ouviu ('Um cliente me disse uma frase esses dias que não saiu da minha cabeça: ...').",
    "NÚMERO/FATO DE MERCADO: use somente um dado que esteja no artigo; se não houver, substitua por uma consequência observável do tema.",
    "CONTRASTE ANTES/DEPOIS: abra mostrando como era antes e como é hoje.",
    "APRENDIZADO PESSOAL: abra com uma confissão/lição sua ('Levei tempo pra entender que...').",
    "PERGUNTA DE ENTRADA: abra com uma pergunta direta e instigante na primeira linha.",
    "BASTIDOR DO TIME: abra por algo que o time (AccessLovers) viveu ou construiu.",
    "OPINIÃO FIRME: abra reto, sem rodeio ('Vou ser direto sobre [tema]:').",
    "CENA DE REUNIÃO/EVENTO: abra numa conversa com clientes/parceiros (sem ser 'de manhã lendo').",
    "ERRO CARO: abra com um erro operacional recorrente e a consequência que ele cria para o caixa ou para o time.",
    "DECISÃO DE CONSELHO: abra pelo tipo de pergunta que uma liderança precisa responder antes de tomar uma decisão financeira.",
    "SINAL FRACO: abra por um pequeno sintoma operacional que costuma revelar um problema maior.",
    "MITO DE MERCADO: abra desmontando uma crença comum sobre o tema, sem criar antagonismo artificial.",
    "CUSTO INVISÍVEL: abra pelo tempo, risco ou previsibilidade que se perde quando a operação continua manual.",
    "PRIMEIRA HORA DO DIA: abra pelo que um CFO ou tesoureiro deveria conseguir enxergar logo cedo.",
    "PERGUNTA DE DIAGNÓSTICO: abra com uma pergunta que ajude o leitor a medir a maturidade da própria operação.",
    "PRINCÍPIO DE CASA: abra por uma convicção de liderança, conectando-a ao tema sem repetir bordões.",
    "MUDANÇA DE ESCALA: abra pelo ponto em que um processo deixa de funcionar porque a empresa cresceu.",
    "EFEITO EM CADEIA: abra mostrando como uma falha pequena se propaga para risco, caixa ou relacionamento.",
    "FUTURO PRÓXIMO: abra pelo que muda na rotina quando a empresa passa a operar com mais visibilidade e integração.",
]

HASHTAG_RULES = [
    ("open finance", ["#OpenFinance", "#IntegraçãoBancária", "#TecnologiaFinanceira"]),
    ("api", ["#APIs", "#IntegraçãoBancária", "#TecnologiaFinanceira"]),
    ("dashboard", ["#CFO", "#AnalyticsFinanceiro", "#DadosEmTempoReal"]),
    ("governança", ["#GovernançaFinanceira", "#Compliance", "#ControlesInternos"]),
    ("dívida", ["#CréditoCorporativo", "#CapitalDeGiro", "#GestãoDeDívidas"]),
    ("crédito", ["#CréditoCorporativo", "#CapitalDeGiro", "#MercadoFinanceiro"]),
    ("cobrança", ["#ContasAReceber", "#Inadimplência", "#FluxoDeCaixa"]),
    ("risco sacado", ["#RiscoSacado", "#CapitalDeGiro", "#SupplyChainFinance"]),
    ("aprova", ["#ContasAPagar", "#AutomaçãoFinanceira", "#Rastreabilidade"]),
    ("conciliação", ["#ConciliaçãoFinanceira", "#ContasAPagar", "#AutomaçãoFinanceira"]),
    ("tesouraria", ["#Tesouraria", "#FluxoDeCaixa", "#GestãoFinanceira"]),
]

def select_hashtags(title, body, limit=5):
    """Mantém marca + tema do artigo, com 3–5 hashtags verificáveis e sem repetição."""
    text = f"{title} {body}".lower()
    tags = ["#Accesstage", "#GestãoFinanceira"]
    for keyword, candidates in HASHTAG_RULES:
        if keyword in text:
            tags.extend(candidates)
    if len(tags) < 5:
        tags.extend(["#TecnologiaFinanceira", "#EficiênciaOperacional", "#CrescimentoSustentável"])
    unique = []
    for tag in tags:
        if tag not in unique:
            unique.append(tag)
    return unique[:limit]

def strip_html(s):
    s = re.sub(r"<(script|style)[^>]*>.*?</\1>", " ", s, flags=re.S | re.I)
    s = re.sub(r"<[^>]+>", " ", s)
    s = html.unescape(s)
    return re.sub(r"\s+", " ", s).strip()

def clean_text(t):
    """Remove travessões e cacoetes; preserva quebras de linha (parágrafos)."""
    out = []
    for line in t.split("\n"):
        # travessão/en-dash/figure-dash -> vírgula
        line = line.replace(" — ", ", ").replace(" – ", ", ").replace("—", ", ").replace("–", ", ").replace("―", ", ")
        line = line.replace(";", ".")                       # sem ponto e vírgula
        line = re.sub(r"[ \t]*,[ \t]*,", ", ", line)        # ", ,"
        line = re.sub(r",[ \t]*\.", ".", line)              # ", ."
        line = re.sub(r"[ \t]+([,.!?:])", r"\1", line)      # espaço antes de pontuação
        line = re.sub(r"[ \t]{2,}", " ", line).strip()      # espaços duplos
        line = re.sub(r"^,\s*", "", line)                   # vírgula órfã no início
        out.append(line)
    return "\n".join(out)

def gen_post(title, body, url, angle, model, emoji_mode="none"):
    emoji_instruction = (
        "Você pode usar no máximo UM emoji, apenas se fizer sentido emocionalmente e nunca em post técnico."
        if emoji_mode == "subtle" else "Não use emojis."
    )
    usr = (f"INPUT — Artigo do blog:\nTÍTULO: {title}\n\nCORPO (resumo textual): {body[:1800]}\n\nURL: {url}\n\n"
           f"ÂNGULO DE ABERTURA OBRIGATÓRIO para este post: {angle}\n"
           f"Lembre: NÃO comece com 'comecei a manhã/dia/semana' nem diga que leu um artigo. "
           f"Escreva com energia e opinião, diferente de um post padrão. {emoji_instruction} "
           f"Não inclua hashtags, pois o sistema as adicionará. Gere agora o post do Celso.")
    for _ in range(3):
        try:
            r = requests.post(URL, headers={"Authorization": f"Bearer {KEY}"},
                json={"model": model, "messages": [{"role":"system","content":SYSTEM},{"role":"user","content":usr}],
                      "temperature": 0.95, "max_tokens": 900}, timeout=90)
            txt = r.json()["choices"][0]["message"]["content"].strip().strip('"').strip()
            if txt and "{URL}" not in txt:
                return txt
        except Exception as e:
            print("   (retry:", str(e)[:50], ")"); time.sleep(1)
    return None

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--csv", default=os.path.join(BASE, "output/articles/lote_veragi_lote2_combined.csv"))
    ap.add_argument("--url_dir", default=os.path.join(BASE, "output/accesstage-site/lote2"))
    ap.add_argument("--url_base", default="https://sowads-accesstage.web.app/lote2")
    ap.add_argument("--out", default=os.path.join(BASE, "output/echo/sowads_echo_lote2.docx"))
    ap.add_argument("--lote", default="Lote 2")
    ap.add_argument("--model", default=MODEL)
    ap.add_argument("--limit", type=int, default=0, help="0 = todos; N = só os N primeiros (amostra)")
    ap.add_argument("--json_out", help="Arquivo JSON com posts para preview e auditoria")
    ap.add_argument("--emoji_mode", choices=["none", "subtle"], default="none",
                    help="none (padrão) ou subtle (no máximo 1 emoji contextual)")
    a = ap.parse_args()
    os.makedirs(os.path.dirname(a.out), exist_ok=True)

    rows = list(csv.DictReader(open(a.csv, encoding="utf-8")))
    if a.limit: rows = rows[:a.limit]
    files = sorted(glob.glob(os.path.join(a.url_dir, "artigo_*.html")))  # ordem = ordem do CSV

    doc = Document()
    st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(11)
    # capa
    h = doc.add_paragraph(); run = h.add_run("Sowads Echo · Posts de Celso Sato"); run.bold = True; run.font.size = Pt(22)
    sub = doc.add_paragraph(); r2 = sub.add_run(f"Emulação do CEO · a partir dos artigos do blog · {a.lote} · {len(rows)} posts")
    r2.font.size = Pt(12); r2.font.color.rgb = RGBColor(0x5b,0x6b,0x7b)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    ok = 0
    generated_posts = []
    for i, row in enumerate(rows):
        title = row.get("post_title", "").strip()
        body = strip_html(row.get("post_content", "") or row.get("meta_description", ""))
        fname = os.path.basename(files[i]) if i < len(files) else ""
        url = f"{a.url_base}/{fname}" if fname else a.url_base
        angle = OPENING_ANGLES[i % len(OPENING_ANGLES)]
        print(f"[{i+1}/{len(rows)}] {title[:46]} | {angle.split(':')[0]}")
        post = gen_post(title, body, url, angle, a.model, a.emoji_mode)
        if not post:
            print("   ✗ falhou"); continue
        post = clean_text(post)
        # O modelo é instruído a não gerar hashtags; removemos qualquer sobra e aplicamos
        # uma seleção rastreável baseada no tema real do artigo.
        post = re.sub(r"(?m)^\s*#[^\n]+\n?", "", post).strip()
        hashtags = select_hashtags(title, body)
        post = f"{post}\n\n{' '.join(hashtags)}"
        generated_posts.append({
            "number": i + 1,
            "title": title,
            "url": url,
            "angle": angle.split(":", 1)[0],
            "copy": post,
            "hashtags": hashtags,
        })
        # rótulo de referência (cinza, pequeno)
        lab = doc.add_paragraph(); lr = lab.add_run(f"POST {i+1} · baseado em: {title}")
        lr.italic = True; lr.font.size = Pt(9); lr.font.color.rgb = RGBColor(0x88,0x88,0x88)
        doc.add_paragraph()
        for para in post.split("\n"):
            if para.strip(): doc.add_paragraph(para.strip())
        ok += 1
        if i < len(rows) - 1:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        time.sleep(0.2)

    doc.save(a.out)
    if a.json_out:
        os.makedirs(os.path.dirname(a.json_out) or ".", exist_ok=True)
        with open(a.json_out, "w", encoding="utf-8") as f:
            json.dump({"persona": "Celso Sato", "lote": a.lote, "posts": generated_posts}, f,
                      ensure_ascii=False, indent=2)
    print(f"\n✓ {ok}/{len(rows)} posts | DOCX: {a.out}")

if __name__ == "__main__":
    main()
